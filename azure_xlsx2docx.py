# Python script to convert Azure Security Benchmark Documentation
# From .xlsx to .docx
# Dependencies: Git, pip gitpython, pip python-docx

from git import Repo #gitpython
from urllib.parse import unquote
import pandas as pd
from pathlib import Path
import os
import shutil
import docx
import re

def azure_xlsx2docx(git_location):

    script_dir = os.path.splitext(__file__)[0]
    git_repo_path = get_local_repo(git_location, script_dir)

    docx_dir = script_dir + "\\Generated DOCX Files"  # separate folder for gen documents
    if os.path.isdir(docx_dir):
        shutil.rmtree(docx_dir)
    os.mkdir(docx_dir)

    xlsx_paths = Path(git_repo_path).glob('*.xlsx')

    for xlsx_path in xlsx_paths:

        filename = os.path.splitext(str(xlsx_path).split("\\")[-1])[0]

        print(f"Loading {filename} ...")
        xlsx = pd.ExcelFile(xlsx_path)
        df_security = pd.read_excel(xlsx, 'Security Profile')
        df_feature = pd.read_excel(xlsx, 'Feature Summary')

        df_feature = df_feature[df_feature['Responsibility'] == 'Customer']

        fill_docx(docx_dir, df_feature, filename)




def fill_docx(docx_dir, df_feature, filename):

    print(f"Generating docx for {filename} ...")
    doc = docx.Document()
    doc.add_heading(filename, 0)

    df_feature.sort_values(by="ASB Control ID", ascending=True, inplace=True)
    df_feature.reset_index(drop=True, inplace=True)

    prev_id = ""
    for i, row in df_feature.iterrows():
        
        if row['ASB Control ID'] != prev_id:
            doc.add_heading(f"{row['ASB Control ID']} - {row['ASB Control Title']}", 3)
        
        p = doc.add_paragraph(style="List Bullet")
        fname_run = p.add_run(f"{row['Feature Name']}")
        fname_run.bold = True
        fname_run = p.add_run(f" - {row['Feature Description']}")
        # p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = 2
        doc.add_paragraph(f"{row['Guidance']}", style="List Bullet 2")

        prev_id = row['ASB Control ID']


    doc.save(os.path.join(docx_dir, filename + ".docx"))

    print(f"Finished\n")
    

    
def get_local_repo(user_input, script_dir):

    repo_dir = script_dir + "\\gitrepo"  # separate folder for git repo

    r1 = r"https://.*"
    if re.match(r1, user_input):
        print("HTTPS repo link found")
        git_folder_https = unquote(user_input)  # link clean up
        r2 = r"((https://github.com/.*?)/tree/.*?/(.*))"
        result = re.search(r2, git_folder_https)

        if result is not None:
            git_repo_https = result.group(2)
            git_folder_path = os.path.join(repo_dir, result.group(3).replace("/", "\\"))

            if not os.path.isdir(script_dir):
                os.mkdir(script_dir)
                os.mkdir(repo_dir)
                print("Cloning repo from Github...")
                Repo.clone_from(git_repo_https, repo_dir)
                print(f"Cloning complete to {repo_dir}")
            else:
                print(f"{script_dir.split('/')[-1]} folder already exists. Continuing using local files...")
        else:
            print(f"Invalid github folder link: {git_folder_https}\nExiting...")  # TODO: fix for no folder repos
            exit(2772)

    elif os.path.isdir(user_input):
        print("Local repo found")
        git_folder_path = user_input


    return git_folder_path


if __name__ == "__main__":
    azure_xlsx2docx("https://github.com/MicrosoftDocs/SecurityBenchmarks/tree/master/Azure%20Offer%20Security%20Baselines/3.0")
    # azure_xlsx2docx("https://github.com/epidemian/snake/tree/test")
    